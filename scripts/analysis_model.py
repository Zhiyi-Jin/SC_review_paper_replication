# -*- coding: utf-8 -*-
"""
This is a script for text analysis on social capital literature.

Created on Fri Mar 14 10:43:58 2025

@author: Zhiyi Jin
"""
# %% Packages
import pandas as pd
import numpy as np
import re
import nltk
#import os
from pathlib import Path

from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from sklearn.feature_extraction.text import CountVectorizer
from nltk import pos_tag
from nltk.corpus import wordnet
from nltk.tokenize import word_tokenize
#from sklearn.model_selection import train_test_split

# Download NLTK resources
nltk.download(['punkt', 'stopwords', 'wordnet'])
nltk.download('punkt_tab')
nltk.download('averaged_perceptron_tagger_eng')

BASE_DIR = Path(__file__).resolve().parent.parent
INPUT_DIR = BASE_DIR / "data" / "processed"
OUTPUT_DIR = BASE_DIR / "results"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Wordcloud
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec

# LDA
from gensim.models import LdaModel
from gensim.matutils import Sparse2Corpus
from gensim.corpora import Dictionary
from gensim import corpora

# Visualization
import pyLDAvis.gensim_models as gensimvis
import pyLDAvis

# Evaluation
from gensim.models import CoherenceModel

# Output
from docx import Document

# Other
#from scipy import sparse
from collections import Counter
from scipy.sparse import hstack
import ast

#import itertools
import requests
from tqdm import tqdm
from gensim.corpora import MmCorpus


# %% Tokenization
# Function to convert POS tags to WordNet format
def get_wordnet_pos(treebank_tag):
    if treebank_tag.startswith('J'):
        return wordnet.ADJ
    elif treebank_tag.startswith('V'):
        return wordnet.VERB
    elif treebank_tag.startswith('N'):
        return wordnet.NOUN
    elif treebank_tag.startswith('R'):
        return wordnet.ADV
    else:
        return wordnet.NOUN  # Default to noun

# Add stopwords manually
custom_stopwords = {"find", "use", "result", "study", "data", "research",
                    "analysis", "effect", "model", "examine", "also",
                    "among", "show", "finding", "factor", "paper", "provide",
                    "survey", "may", "include", "suggest", "focus", "method",
                    "paper", "outcome", "however", "relate", "need", "significant",
                    "develop", "need", "level", "influence", "impact", "relationship",
                    "base", "explore", "theory", "change", "well", "within",
                    "article", "positive", "affect", "aim", "association",
                    "test", "implication", "conduct", "variable", "literature",
                    "evidence", "explain", "significantly", "likely", "whether",
                    "lead", "author", "analyze", "ie", "become", "occupation",
                    "social", "capital", "network", "labor", "labour", "market",
                    "different", "people", "associate", "occupational", "low",
                    "high", "measure", "one", "two", "behavior", "important",
                    "new", "investigate", "human", "difference", "sample",
                    "approach", "less", "first", "indicate", "type", "conclusion",
                    "good", "improve", "report", "number", "make", "right", "great",
                    "employee", "employment", "worker", "consider", "across",
                    "contribute", "many", "perceive", "collect", "importance",
                    "elsevier", "draw", "reduce", "understand", "lack", "de",
                    "take", "objective", "demonstrate", "propose", "current",
                    "analyse", "regard", "ci", "highlight", "small", "seek",
                    "set", "various", "exist", "issue", "obtain", "often", "receive",
                    "apply", "thus", "know", "little", "previous", "could",
                    "recent", "follow", "give", "predict", "particularly", "yet",
                    "eg", "via", "rather", "whereas", "ltd", "major", "leave",
                    "several", "maintain", "furthermore", "establish", "consequence",
                    "likelihood", "complete", "least", "odds", "critical", "process",
                    "three", "large", "compare", "strong", "reveal"}  

try:
    english_stopwords = stopwords.words('english')
except NameError:
    from nltk.corpus import stopwords as nltk_stopwords
    english_stopwords = nltk_stopwords.words('english')

stop_words = set(english_stopwords).union(custom_stopwords)

# Semantic generalization
# Define a semantic mapping
semantic_map = {
    "woman": "gender", 
    "young":"age",
    "member":"membership", 
    "structural":"structure",
    "educational":"education",
    "old":"age", 
    "resident":"residence",
    "youth":"age",
    "female":"gender",
    "male":"gender",
    "leader":"leadership",
    "chinese":"china",
    "womens":"gender",
    "sex":"gender",
    "depend":"dependence",
    "diverse":"diversity",
    "sustainable":"sustainability",
    "entrepreneurial" : "entrepreneurship",
    "adjust" : "adjustment",
    "residential" : "residence",
    "independent" : "independence",
    "mobile" : "mobility",
    "produce" : "production",
    "entrepreneur" : "entrepreneurship",
    "reference" : "referral",
    "directly" : "direct",
    "racial" : "race",
    "agent" : "agency",
    "childrens" : "child",
    "skilled" : "skill",
    "religious" : "religion",
    "dependent" : "dependence",
    "inform" : "information",
    "organisation" : "organization",
    "enter" : "entry",
    "marital" : "marriage",
    "stay" : "remain",
    "competitive" : "competition",
    "vulnerable" : "vulnerability",
    "secure" : "security",
    "parental" : "parent",
    "organisational" : "organization",
    "organizational": "organization",
    "gendered" : "gender",
    "ethnicity" : "ethnic",
    "employed" : "employ",
    "returnees" : "migrant",
    "immigration" : "migrant",
    "immigrant": "migrant",
    "migrate" : "migrant",
    "migration": "migrant",
    "protective" : "protection",
    "educate" : "education",
    "productive" : "productivity",
    "independently" : "independence",
    "indirectly" : "indirect",
    "communicate" : "communication",
    "compete" : "competition",
    "understood" : "understanding",
    "childhood" : "child",
    "coethnic" : "ethnic",
    "invest" : "investment",
    "middleaged" : "age",
    "efficient" : "efficiency",
    "compensate" : "compensation",
    "safe" : "safety",
    "reside" : "residence",
    "desirable" : "desire",
    "lowskilled" : "skill"
}

def generalize_tokens(tokens):
    return [semantic_map.get(token, token) for token in tokens]

# Preprocessing function
def preprocess_text(text):
    if not isinstance(text, str):
        return []

    text = text.lower()
    text = re.sub(r'[^a-zA-Z\s]', '', text)
    tokens = word_tokenize(text)
    
    # Initialize lemmatizer
    lemmatizer = WordNetLemmatizer()
    
    # POS tagging
    pos_tags = pos_tag(tokens)
    
    # Lemmatize with POS tags
    lemmatized_tokens = []
    for word, tag in pos_tags:
        pos = get_wordnet_pos(tag)
        lemma = lemmatizer.lemmatize(word, pos=pos)
        lemmatized_tokens.append(lemma)
    
    # Remove stopwords AFTER lemmatization
    tokens = [word for word in lemmatized_tokens if word not in stop_words]
    
    # Apply semantic generalization
    generalized_tokens = generalize_tokens(tokens)
    
    return generalized_tokens  

# Load data
try:
    pd
except NameError:
    import pandas as pd

df = pd.read_csv(INPUT_DIR / "articles_clean.csv", encoding="latin1")

df['processed_abstract'] = df['abstract'].apply(preprocess_text)
df['processed_abstract'] = df['processed_abstract'].apply(lambda tokens: ' '.join(tokens))

total_tokens = df['processed_abstract'].apply(lambda text: len(text.split())).sum()
print(f"Total number of tokens: {total_tokens}")

# Function to extract words from word documents
def extract_highlighted_words(docx_path):
    docx_path = Path(docx_path)
    if not docx_path.exists():
        print(f"Warning: file not found, skipping {docx_path}")
        return set()

    doc = Document(docx_path)
    highlighted_words = set()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.font.highlight_color is not None:
                            text = run.text.strip()
                            if text:
                                highlighted_words.add(text)
    return highlighted_words

# Extract colored words from both files
word_zj = extract_highlighted_words(INPUT_DIR / "term_frequency_update_ZJ.docx")
word_md = extract_highlighted_words(INPUT_DIR / "term_frequency_update_MvD.docx")
word_cs = extract_highlighted_words(INPUT_DIR / "term_frequency_update_CS.docx")

# Find words with a frequency of 2 or higher
all_words = list(word_zj) + list(word_md) + list(word_cs)
word_counts = Counter(all_words)
remo_words = [word for word, count in word_counts.items() if count >= 2]

# Create document-term matrix
vectorizer = CountVectorizer(max_df=0.80, # Remove terms in >80% of documents
                                          # same as 0.99
                             min_df=0.005,    # Remove terms in <0.5% documents
                             stop_words=list(remo_words) # Add more stopwords
                             )

doc_term_matrix = vectorizer.fit_transform(df['processed_abstract'])

# %% Check on the document-term matrix
# Get feature names (terms)
terms = vectorizer.get_feature_names_out()

# Convert to DataFrame (rows = documents, columns = terms)
doc_term_df = pd.DataFrame(
    doc_term_matrix.toarray(),
    columns=terms
)

# Display the first few rows
print(doc_term_df.head())

# Sum term frequencies across all documents
term_freq = doc_term_matrix.sum(axis=0).A1  # Convert to 1D array

# Create a DataFrame of terms and their frequencies
term_freq_df = pd.DataFrame({
    'term': terms,
    'frequency': term_freq
})

# Sort by frequency
term_freq_df = term_freq_df.sort_values('frequency', ascending=False)

# Top 10 terms
print(term_freq_df.head(20))

# Output the term frequency table
#doc = Document()
#doc.add_heading('Term Frequency Table', level=1)

# Create table with header row
#table = doc.add_table(rows=1, cols=len(term_freq_df.columns))
#table.style = 'Table Grid'

# Add header
#header_cells = table.rows[0].cells
#for i, column_name in enumerate(term_freq_df.columns):
#    header_cells[i].text = column_name

# Add data rows
#for index, row in term_freq_df.iterrows():
#    row_cells = table.add_row().cells
#    for i, value in enumerate(row):
#        row_cells[i].text = str(value)

# Save the document
#doc.save('term_frequency.docx')

# Convert term frequencies to a dictionary
term_freq_dict = dict(zip(term_freq_df['term'], term_freq_df['frequency']))

# Generate the word cloud
wordcloud = WordCloud(
    width=800,
    height=400,
    background_color='white',
    max_words=100,  # Limit to top 100 terms
).generate_from_frequencies(term_freq_dict)

# Display the word cloud
plt.figure(figsize=(10, 5))
plt.imshow(wordcloud, interpolation='bilinear')
plt.axis('off')  # Hide axes
plt.title("Top Terms in Corpus", fontsize=16)
plt.show()

#%% Abstract model
# Convert the document-term matrix to a Gensim corpus
corpus = Sparse2Corpus(doc_term_matrix, documents_columns=False)

# Create a dictionary
dictionary = Dictionary.from_corpus(corpus, 
                                    id2word=dict(zip(range(len(terms)), terms)))

print('Number of unique tokens: %d' % len(dictionary))
print('Number of documents: %d' % len(corpus))

# Model evaluation
# Evaluate cv coherence for different k values 
# Prepare tokenized documents for coherence calculation
tokenized_docs = df['processed_abstract'].str.split().tolist()

def train_lda_models(corpus, dictionary, tokenized_docs, 
                     k_range, passes=15, random_state=42):
    cv_scores, umass_scores, models = [], [], []

    for k in k_range:
        print(f"Training LDA for {k} topics...")
        lda = LdaModel(corpus=corpus, num_topics=k, 
                       id2word=dictionary, passes=passes, 
                       random_state=random_state)
        models.append(lda)
        cv = CoherenceModel(model=lda, texts=tokenized_docs, 
                            dictionary=dictionary, 
                            coherence='c_v', topn=15).get_coherence()
        umass = CoherenceModel(model=lda, corpus=corpus, 
                               dictionary=dictionary, 
                               coherence='u_mass', topn=15).get_coherence()
        cv_scores.append((k, cv))
        umass_scores.append((k, umass))
    return models, cv_scores, umass_scores

def plot_scores(scores, title, ylabel):
    ks, vals = zip(*scores)
    plt.figure(figsize=(10, 5))
    plt.plot(ks, vals, marker='o')
    plt.xlabel("Number of Topics")
    plt.ylabel(ylabel)
    plt.title(title)
    plt.grid(True)
    plt.show()

def calculate_diversity(model, dictionary, topn=10):
    all_words = [dictionary[word_id] 
                 for topic_id in range(model.num_topics)
                 for word_id, _ in model.get_topic_terms(topic_id, topn=topn)]
    return len(set(all_words)) / len(all_words) * 100 if all_words else 0

def plot_diversity(models, dictionary, topn=15):
    scores = [(model.num_topics, 
               calculate_diversity(model, dictionary, topn)) for model in models]
    plot_scores(scores, "Topic Diversity", "Diversity Score (%)")

def plot_perplexity(models, corpus):

    log_perps = [model.log_perplexity(corpus) for model in models]
    true_perps = [np.exp(-lp) for lp in log_perps]
    ks = [model.num_topics for model in models]

    plt.figure(figsize=(10, 6))
    plt.plot(ks, true_perps, marker='o')
    plt.title('LDA Model Perplexity by Number of Topics')
    plt.xlabel('Number of Topics')
    plt.ylabel('Perplexity')
    plt.grid(True)
    plt.tight_layout()
    plt.show()

# === Run All Steps ===
models, cv_scores, umass_scores = train_lda_models(corpus, dictionary, 
                                                   tokenized_docs,
                                                   range(2,31))

plot_scores(cv_scores, "CV Coherence by TNumber of Topics", "Coherence Score")
plot_scores(umass_scores, "UMass Coherence by Number of Topics", "Coherence Score")
plot_diversity(models, dictionary, topn=15)
plot_perplexity(models, corpus) #FULL DATA

# save the combined corpus
MmCorpus.serialize(str(OUTPUT_DIR / "abs_model" / "ab_corpus.mm"), corpus) 
# save the combined dictionary
dictionary.save(str(OUTPUT_DIR / "abs_model" / "ab_dictionary.dict"))

# Visualize optimal models
for i in range(0, 28):
    optimal_model = models[i]
    vis_data = gensimvis.prepare(optimal_model, corpus, dictionary)
    pyLDAvis.save_html(vis_data, str(OUTPUT_DIR / "abs_model" / f'lda_model_update_{i+1}.html'))


#%% Test filter more terms
# vectorizer_1 = CountVectorizer(max_df=0.80, # 80% = 99%
#                              min_df=0.001   # 0.5%
#                              ) #prune
# doc_term_matrix_1 = vectorizer_1.fit_transform(df['processed_abstract'])

# # Check if any documents with all 0s
# if sparse.issparse(doc_term_matrix_1):
#     zero_rows = np.where(doc_term_matrix_1.getnnz(axis=1) == 0)[0]
#     if len(zero_rows) > 0:
#         print(f"Found {len(zero_rows)} all-zero documents at positions: {zero_rows}")
#     else:
#         print("No documents with all zeros found")

# # Get feature names (terms)
# terms_1 = vectorizer_1.get_feature_names_out()

# # Convert to DataFrame (rows = documents, columns = terms)
# doc_term_df_1 = pd.DataFrame(
#     doc_term_matrix_1.toarray(),
#     columns=terms_1
# )


# # Sum term frequencies across all documents
# term_freq_1 = doc_term_matrix_1.sum(axis=0).A1  # Convert to 1D array

# # Create a DataFrame of terms and their frequencies
# term_freq_df_1 = pd.DataFrame({
#     'term': terms_1,
#     'frequency': term_freq_1
# })

# # Sort by frequency
# term_freq_df_1 = term_freq_df_1.sort_values('frequency', ascending=False)

# # Top 10 terms
# print(term_freq_df_1.head(20))

# # Output the term frequency table
# doc_1 = Document()
# doc_1.add_heading('Term Frequency Table (update)', level=1)

# # Create table with header row
# table_1 = doc_1.add_table(rows=1, cols=len(term_freq_df_1.columns))
# table_1.style = 'Table Grid'

# Add header
# header_cells_1 = table_1.rows[0].cells
# for i, column_name in enumerate(term_freq_df_1.columns):
#     header_cells_1[i].text = column_name

# Add data rows
# for index, row in term_freq_df_1.iterrows():
#     row_cells = table_1.add_row().cells
#     for i, value in enumerate(row):
#         row_cells[i].text = str(value)

# Save the document
# doc_1.save('term_frequency_update.docx')

#%% Tokenize the reference into BOW
# ===== Extact DOIs ======
def extract_dois(reference_str):
    dois = []
    # Check if the input is a string (skip NaN/float entries)
    if isinstance(reference_str, str):
        references = reference_str.split(' and ')
        for ref in references:
            # Find all DOIs in the reference
            matches = re.findall(r'DOI[\s\[]*([^,\s\]]+)', ref)
            dois.extend(matches)
    return dois

# Apply the function to the cited references column to create doi_tokens
df['doi_tokens'] = df['CR'].apply(extract_dois)

# Count document frequency for raw DOIs
doi_doc_freq = Counter()

for doi_list in df['doi_tokens']:
    unique_dois = set(doi_list)
    doi_doc_freq.update(unique_dois)

# Filter DOIs based on frequency thresholds
num_docs = len(df)
upper_thresh = 0.80 * num_docs
lower_thresh = 0.005 * num_docs

valid_dois = {doi for doi, freq in doi_doc_freq.items() if lower_thresh < freq < upper_thresh}
invalid_dois = {"DOI", "NG"}
valid_dois = valid_dois - invalid_dois

# Clean the dois
def clean_doi(doi):
    """
    Cleans a DOI string by:
    - Removing leading 'https://doi.org/' or 'http://doi.org/'
    - Replacing any double slashes after the prefix (like '10.xxxx//') with a single slash
    - Returning a valid DOI format (10.xxxx/xxxx)
    """
    if not isinstance(doi, str):
        return None

    # Step 1: Remove leading URL
    doi = doi.strip()
    doi = re.sub(r'^https?://doi\.org/', '', doi)

    # Step 2: Replace double slashes immediately after the '10.xxxx' prefix
    doi = re.sub(r'^(10\.\d{4,9})//', r'\1/', doi)

    # Step 3: Match and return the correct DOI format
    match = re.match(r'^10\.\d{4,9}/\S+$', doi)
    if match:
        return match.group(0)
    else:
        return None

# Get cleaned dois to match through API
cleaned_dois = set()
    
for doi in valid_dois:
    cleaned = clean_doi(doi)
    if cleaned:
        cleaned_dois.add(cleaned)

# Match DOIs to author and year
# create a dictionary
doi_to_triple = {}

# Match with the author and year 
def get_metadata_from_doi(doi):
    url = f"https://api.crossref.org/works/{doi}"
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        data = response.json()
        authors = data['message'].get('author', [])
        first_author = authors[0]['family'] if authors else 'Unknown'
        year = data['message']['issued']['date-parts'][0][0]
        return f"{first_author},{year},{doi}"
    except Exception as e:
        print(f"Failed to fetch metadata for DOI {doi}: {e}")
        return f"Unknown,None,{doi}"  
    
for doi in tqdm(cleaned_dois, desc="Fetching DOI metadata"):
    doi_to_triple[doi] = get_metadata_from_doi(doi)

# Manually add those are not fetched successfully
doi_to_triple['10.1016/jj.deveco.2004.11.006'] = 'Wahba,2005,10.1016/j.jdeveco.2004.11.006'
doi_to_triple['10.4324/9780429306419-102/american-apartheid-segregation-making-underclass-douglas-massey-nancy-denton'] = 'Massey,1993,10.4324/9780429306419-102'
doi_to_triple['10.2307/20159100'] = 'Inkpen,2005,10.2307/20159100'
doi_to_triple['10.2307/20159747'] = 'Perry‑Smith,2006,10.2307/20159747'

# Convert them into a tri-gram in a column
df['doi_trigrams'] = df['doi_tokens'].apply(
    lambda dois: [doi_to_triple[doi] for doi in dois if doi in doi_to_triple]
)

# Print the IDs of rows where doi_trigrams is empty
empty_doi_ids = df[df['doi_trigrams'].apply(lambda x: len(x) == 0)]['ID']
print(empty_doi_ids.tolist())
df['doi_trigrams'].apply(lambda x: len(x) == 0).sum()

# store the data with cleaned dois
df.to_csv(OUTPUT_DIR / "cleaned_doi_data_2.csv", index=False)

#%% Check those articles without CR after filter

# Filter again to ensure documents have tokens
df = df[df['doi_trigrams'].apply(len) > 0]

# Check the abstract topic distribution among those articles without references
# Create a matrix where rows = documents, columns = topics
all_topics = [optimal_model.get_document_topics(doc_bow) for doc_bow in corpus]

# print topic distribution for first 5 documents
for i, doc_topics in enumerate(all_topics[:5]):
    print(f"Document {i}: {doc_topics}")
    
topic_matrix = np.zeros((len(corpus), num_topics))

for doc_idx, doc_topics in enumerate(all_topics):
    for topic_id, prob in doc_topics:
        topic_matrix[doc_idx, topic_id] = prob

# Step 1: Prepare list of document indices from empty_doi_ids
doc_indices = empty_doi_ids.index.tolist()  # Adjust if needed

# Step 2: Get topic distributions for these 843 documents
num_topics = optimal_model.num_topics
topic_matrix_subset = np.zeros((len(doc_indices), num_topics))

for i, doc_idx in enumerate(doc_indices):
    doc_topics = optimal_model.get_document_topics(corpus[doc_idx])
    for topic_id, prob in doc_topics:
        topic_matrix_subset[i, topic_id] = prob

# Step 3: Sum across documents to get total contribution of each topic
topic_sums = topic_matrix_subset.sum(axis=0)

# Step 4: Plot histogram
plt.figure(figsize=(10, 6))
plt.bar(range(num_topics), topic_sums, color='skyblue')
plt.xlabel("Topic ID")
plt.ylabel("Total Topic Probability (across 843 documents)")
plt.title("Topic Distribution in Documents with Empty DOI IDs")
plt.xticks(range(num_topics))
plt.grid(axis='y', linestyle='--', alpha=0.7)
plt.tight_layout()

plt.savefig(OUTPUT_DIR / "empty_doi_dist.jpg", format='jpg', dpi=300, bbox_inches='tight')
plt.close()

# Make a wordcloud for the documents
subset_df = doc_term_df.loc[empty_doi_ids.index]

# Sum the term frequencies across these documents
term_frequencies = subset_df.sum(axis=0)

# Convert to dictionary for WordCloud
term_freq_dict = term_frequencies.to_dict()

# Create and plot the word cloud
wordcloud = WordCloud(width=1000, height=500, background_color='white', colormap='tab10')
wordcloud.generate_from_frequencies(term_freq_dict)

plt.figure(figsize=(14, 7))
plt.imshow(wordcloud, interpolation='bilinear')
plt.axis('off')
plt.title("Word Cloud for 843 Documents with Empty DOI IDs", fontsize=16)
plt.show()


#%% Reference model
# Create dictionary and corpus for reference-LDA
# Extract documents: each row is a list of trigram tokens
reference_documents = df['doi_trigrams'].tolist()

reference_dictionary = corpora.Dictionary(reference_documents)
reference_corpus = [reference_dictionary.doc2bow(doc) for doc in reference_documents]
reference_tokenized_documents = [
    [reference_dictionary.token2id[token] for token in doc if token in dictionary.token2id]
    for doc in reference_documents]

print('Number of unique tokens:', len(reference_dictionary))
print('Number of documents:', len(reference_corpus))

# Train LDA model
r_models, r_cv_scores, r_umass_scores = train_lda_models(
    reference_corpus, reference_dictionary, reference_documents, 
    k_range=range(2, 31))

plot_scores(r_cv_scores, "Reference model: CV Coherence by Number of Topics", 
            "Coherence Score")
plot_scores(r_umass_scores, "Reference model: UMass Coherence by Number of Topics", 
            "Coherence Score")
plot_diversity(r_models, reference_dictionary, topn=15)
plot_perplexity(r_models, reference_corpus)
    
# Visulize the optimal result
# n = 5
num_topics_5 = r_models[3].num_topics
topn = 10
for topic_id in range(num_topics_5):
    topic_terms = r_models[3].show_topic(topic_id, topn=10)
    terms, weights = zip(*topic_terms)

    plt.figure(figsize=(10, 5))
    plt.barh(terms, weights, color='skyblue')
    plt.gca().invert_yaxis()  # Highest at top
    plt.title(f"Top 10 Tokens for Topic {topic_id}")
    plt.xlabel("Weight")
    plt.tight_layout()
    plt.show()

# n = 6
num_topics_6 = r_models[4].num_topics
for topic_id in range(num_topics_6):
    topic_terms = r_models[4].show_topic(topic_id, topn=10)
    terms, weights = zip(*topic_terms)

    plt.figure(figsize=(10, 5))
    plt.barh(terms, weights, color='skyblue')
    plt.gca().invert_yaxis()  # Highest at top
    plt.title(f"Top 10 Tokens for Topic {topic_id}")
    plt.xlabel("Weight")
    plt.tight_layout()
    plt.show()


#%% Abstract-Reference model
file_path = OUTPUT_DIR / "cleaned_doi_data_2.csv"
df_doi = pd.read_csv(file_path)

# Vectorize abstract with original filters
abstract_vectorizer = CountVectorizer(
    max_df=0.80,
    min_df=0.005,
    stop_words=list(remo_words))
abstract_matrix = abstract_vectorizer.fit_transform(df['processed_abstract'])
abstract_terms = abstract_vectorizer.get_feature_names_out()

# Create trigram strings from the DOI dataframe
df['doi_trigrams'] = df_doi['doi_trigrams'].apply(
    lambda x: ast.literal_eval(x) if isinstance(x, str) else x
)

df['trigram_str'] = df['doi_trigrams'].apply(
    lambda x: ' '.join(x) if isinstance(x, list) and x else ''
)

# Vectorize trigrams with error handling
trigram_vectorizer = CountVectorizer(
    min_df=1,
    max_df=1.0,
    tokenizer=lambda x: x.split(),
    preprocessor=lambda x: x if isinstance(x, str) else ''
)
trigram_matrix = trigram_vectorizer.fit_transform(df['trigram_str'])
trigram_terms = trigram_vectorizer.get_feature_names_out()

# Combine matrices
combine_doc_term_matrix = hstack([abstract_matrix, trigram_matrix], format='csr')
combine_terms = np.concatenate([abstract_terms, trigram_terms])

combine_corpus = Sparse2Corpus(combine_doc_term_matrix, documents_columns=False)
combine_dictionary = Dictionary.from_corpus(
    combine_corpus, 
    id2word=dict(zip(range(len(combine_terms)), combine_terms))
)

# Build tokenized docs
tokenized_docs = []

for doc_row in combine_doc_term_matrix:
    token_ids = doc_row.indices  # non-zero term indices
    tokens = [combine_dictionary[id] for id in token_ids]
    tokenized_docs.append(tokens)


# Run LDA analysis
c_models, c_cv_scores, c_umass_scores = train_lda_models(
    combine_corpus, combine_dictionary, tokenized_docs, k_range=range(2, 31))

plot_scores(c_cv_scores, "CV Coherence by Number of Topics", "Coherence Score")
plot_scores(c_umass_scores, "UMass Coherence by Number of Topics", "Coherence Score")
plot_diversity(c_models, combine_dictionary, topn=15)
plot_perplexity(c_models, combine_corpus) #FULL DATA

# save the result
for i, model in enumerate(c_models):
    model.save(str(OUTPUT_DIR / "abs_ref model" / f'model_{i}.model'))

# save the combined corpus
MmCorpus.serialize(str(OUTPUT_DIR / "abs_ref model"/ "combine_corpus.mm"), combine_corpus) 

# save the combined dictionary
combine_dictionary.save(str(OUTPUT_DIR / "abs_ref model" / "combine_dictionary.dict"))

# load the result
loaded_models = []
for i in range(1):  # or however many models you saved
    loaded_models.append(LdaModel.load(str(OUTPUT_DIR / "abs_ref model" / f'model_{i}.model')))

#%% Check the No.6 topic (rural-urban)

# Get topic distribution for each document
doc_topic_probs = []
for i, doc_bow in enumerate(combine_corpus):
    topics = optimal_c_model.get_document_topics(doc_bow, minimum_probability=0.0)
    topic6_prob = dict(topics).get(6, 0.0)  # Topic index 6 (i.e., 7th topic in Gensim)
    doc_topic_probs.append((i, topic6_prob))

# Sort by topic 6 weight descending
top_docs = sorted(doc_topic_probs, key=lambda x: x[1], reverse=True)[:20]

# Show metadata for top documents
top_doc_indices = [doc_id for doc_id, _ in top_docs]
top_df = df.iloc[top_doc_indices][['ID', 'title', 'abstract', 'year', 'doi']].copy()

# Also include topic 6 weights
top_df['topic_6_weight'] = [prob for _, prob in top_docs]
top_df.to_csv(OUTPUT_DIR / "top_rural_urban.csv", index=False)

# Extract only the topic 6 probabilities
topic6_weights = [prob for _, prob in doc_topic_probs]

# Plot histogram
plt.figure(figsize=(10, 6))
plt.hist(topic6_weights, bins=20, edgecolor='black')
plt.title('Distribution of Topic 6 Across Documents')
plt.xlabel('Topic 6 weight in document')
plt.ylabel('Number of Documents')
plt.grid(True)
plt.tight_layout()

#%% Extract books from the CR
file_path = OUTPUT_DIR / "cleaned_doi_data_2.csv"
df_doi = pd.read_csv(file_path)

# Canonical books with matching patterns
books = {
    "Bourdieu, P., 1986": r"BOURDIEU.*1986.*(FORMS OF CAPITAL|HANDBOOK OF THEORY|RICHARDSON)",
    
    "Burt, R. S., 1992": r"BURT.*1992.*STRUCTURAL HOLES",
    
    "Coleman, J. S., 1990": r"COLEMAN.*1990.*FOUNDATIONS OF SOCIAL THEORY",
    
    "Putnam, R. D., 2000": r"PUTNAM.*2000.*BOWLING ALONE",
    
    "Putnam, R. D., 1993": r"PUTNAM.*1993.*MAKING DEMOCRACY WORK",
    
    "Lin, N., 2001": r"LIN.*2001.*SOCIAL CAPITAL"
}

def extract_books_precise(cr):
    if pd.isna(cr):
        return []
    cr = cr.upper()
    return [
        book for book, pattern in books.items()
        if re.search(pattern, cr)
    ]

# Apply extraction
df_doi["cr_book"] = df_doi["CR"].apply(extract_books_precise)

# Explode for corpus-level counting
exploded = df_doi.explode("cr_book").dropna(subset=["cr_book"])

# Global frequency table
freq_table = (
    exploded["cr_book"]
    .value_counts()
    .reset_index()
    .rename(columns={"index": "book", "cr_book": "frequency"})
)

freq_table























